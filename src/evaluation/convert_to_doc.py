from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from collections import defaultdict
import json
import ast

with open("eval_combined_clean.json", encoding="utf-8") as f:
    data = json.load(f)

questions = data.get("perquestionresults", [])

DOC_LABELS = {
    "Adhesive_capsulitis_JOSPT.pdf":
        "Adhesive Capsulitis",
    "AdhesivecapsulitisJOSPT.pdf":
        "Adhesive Capsulitis",
    "Copy of Skulder og skulderbue.pdf":
        "Shoulder Examination and Pathology",
    "Rotator cuff tendinopathy CPG.pdf":
        "Rotator Cuff Tendinopathy",
    "Shoulderdoc - Shoulder Rehab Book.pdf":
        "Rehabilitation and Exercise",
    "Subacromial pain syndrome.pdf":
        "Subacromial Pain Syndrome",
    "Therapeutic_exercise_Foundations_and_techniques_by_Colby_Lynn_Allen-572-650.pdf":
        "Therapeutic Exercise and Postoperative Care",
    "TherapeuticexerciseFoundationsandtechniquesbyColbyLynnAllen-572-650.pdf":
        "Therapeutic Exercise and Postoperative Care",
}

def get_options(q):
    options = q.get("options", {})
    if isinstance(options, str):
        try:
            options = ast.literal_eval(options)
        except Exception:
            options = {}
    return options

def parse_gold_answer(q):
    # Try both key names
    gold = q.get("gold_answer") or q.get("goldanswer", "")

    # Dict format
    if isinstance(gold, dict):
        letter = str(gold.get("Answer", "")).strip().lower().rstrip(".")
        text   = gold.get("Text_Answer", "").strip()
        return letter, text

    # String representation of dict
    gold_str = str(gold).strip()
    if not gold_str:
        return "", ""

    # Try parsing as dict string
    if gold_str.startswith("{"):
        try:
            gold_dict = ast.literal_eval(gold_str)
            letter = str(gold_dict.get("Answer", "")).strip().lower().rstrip(".")
            text   = gold_dict.get("Text_Answer", "").strip()
            return letter, text
        except Exception:
            pass

    # Plain string format: "Answer e, TextAnswer some text"
    import re
    letter = ""
    text   = ""

    letter_match = re.search(r'Answer\s+([a-eA-E])', gold_str)
    if letter_match:
        letter = letter_match.group(1).lower()

    if "TextAnswer" in gold_str:
        text = gold_str.split("TextAnswer")[-1].strip()
    elif "Text_Answer" in gold_str:
        text = gold_str.split("Text_Answer")[-1].strip().lstrip(":").strip()

    return letter, text


doc = Document()

title = doc.add_heading(
    "Shoulder MCQ Question Bank — Final Curated Set", 0
)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph(
    f"Total questions: {len(questions)}  |  "
    f"Source: Multi-model generation (Qwen3-14B, Qwen3-8B, Mistral-7B)  |  "
    f"Status: Ready for peer review"
)
doc.add_paragraph("")

by_doc = defaultdict(list)
for q in questions:
    doc_name = q.get("document", "Unknown")
    by_doc[doc_name].append(q)

q_number = 1

for doc_file, qs in by_doc.items():
    label = DOC_LABELS.get(doc_file, doc_file)
    doc.add_heading(label, level=1)

    for q in qs:
        question_text = q.get("question", "")
        options       = get_options(q)

        # Question stem
        p   = doc.add_paragraph(style="Normal")
        run = p.add_run(f"Q{q_number}. {question_text}")
        run.bold      = True
        run.font.size = Pt(11)

        # Answer options
        if isinstance(options, dict):
            for letter, text in sorted(options.items()):
                doc.add_paragraph(
                    f"    {letter}) {text}",
                    style="List Bullet"
                )

        # Correct answer
        letter, text = parse_gold_answer(q)

        # Fallback: get text from options dict if text is empty
        if not text and letter and isinstance(options, dict):
            text = options.get(letter, options.get(letter.upper(), ""))

        ans_para = doc.add_paragraph(style="Normal")
        ans_run  = ans_para.add_run(
            f"Correct answer: {letter.upper()}) {text}"
        )
        ans_run.bold           = True
        ans_run.font.color.rgb = RGBColor(0x1F, 0x7A, 0x1F)
        ans_run.font.size      = Pt(11)

        doc.add_paragraph("")
        q_number += 1

output = "Shoulder_MCQ_Question_Bank.docx"
doc.save(output)
print(f"Saved: {output}  ({q_number - 1} questions)")